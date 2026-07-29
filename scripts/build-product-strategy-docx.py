from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PACKAGE = ROOT / "deliverables" / "LocalAkademi-Product-Strategy-Package"
SOURCE = PACKAGE / "02-Product-Strategy.md"
OUTPUT = PACKAGE / "02-Product-Strategy.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "667085"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True, color=INK)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, color=INK)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])
    set_run_font(run, size=9, color=MUTED)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    if "Lead Callout" not in styles:
        callout = styles.add_style("Lead Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout = styles["Lead Callout"]
    callout.font.name = "Calibri"
    callout.font.size = Pt(11.5)
    callout.font.bold = True
    callout.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(8)
    callout.paragraph_format.space_after = Pt(12)
    callout.paragraph_format.line_spacing = 1.15


def add_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"
    widths = [9360 // cols] * cols
    widths[-1] += 9360 - sum(widths)
    for ridx, values in enumerate(rows):
        cells = table.add_row().cells
        for cidx, value in enumerate(values):
            p = cells[cidx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_inline(p, value)
            if ridx == 0:
                set_cell_shading(cells[cidx], LIGHT)
                for run in p.runs:
                    run.bold = True
            cells[cidx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def markdown_body(doc: Document, lines: list[str]) -> None:
    i = 0
    table_rows: list[list[str]] = []
    while i < len(lines):
        raw = lines[i].rstrip()
        stripped = raw.strip()
        if not stripped:
            if table_rows:
                add_table(doc, table_rows)
                table_rows = []
            i += 1
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [cell.strip() for cell in stripped.strip("|").split("|")]
            if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                i += 1
                continue
            table_rows.append(cells)
            i += 1
            continue
        if table_rows:
            add_table(doc, table_rows)
            table_rows = []
        if stripped.startswith("### "):
            doc.add_paragraph(stripped[4:], style="Heading 2")
        elif stripped.startswith("## "):
            doc.add_paragraph(stripped[3:], style="Heading 1")
        elif stripped.startswith("# "):
            doc.add_paragraph(stripped[2:], style="Heading 1")
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Lead Callout")
            add_inline(p, stripped[2:])
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), PALE_BLUE)
            p_pr.append(shd)
        elif re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, stripped)
            p.paragraph_format.widow_control = True
        i += 1
    if table_rows:
        add_table(doc, table_rows)


def main() -> None:
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    body_start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))

    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    setup_styles(doc)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = hp.add_run("LOCALAKADEMİ  |  ÜRÜN STRATEJİSİ")
    set_run_font(hr, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("LA-MP-02  •  27 Temmuz 2026  •  Sayfa ")
    set_run_font(fr, size=9, color=MUTED)
    add_page_field(fp)

    doc.add_paragraph().paragraph_format.space_after = Pt(22)
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    kr = kicker.add_run("STRATEJİ NOTU  •  V1.0")
    set_run_font(kr, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    tr = title.add_run("LocalAkademi Ürün Stratejisi")
    set_run_font(tr, size=25, color="111827", bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    sr = subtitle.add_run("Kontrollü beta, içerik derinliği ve güvenilir AI karar çerçevesi")
    set_run_font(sr, size=13.5, color=MUTED)

    for label, value in [
        ("Belge kodu", "LA-MP-02"),
        ("Sürüm", "1.0"),
        ("Tarih", "27 Temmuz 2026"),
        ("Durum", "Uygulama ve beta kararları için ana strateji"),
        ("Kapsam", "Türkiye pazarı, LocalAkademi v1.0 ve kontrollü beta"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        lr = p.add_run(f"{label}: ")
        set_run_font(lr, size=10.5, color=INK, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5, color=INK)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    callout = doc.add_paragraph(style="Lead Callout")
    add_inline(callout, "Ana karar: Yeni özellik sayısını artırmadan önce güvenlik, içerik derinliği, ilk değer anı ve kontrollü beta sonuçları kanıtlanmalıdır.")
    callout_pr = callout._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), PALE_BLUE)
    callout_pr.append(shd)

    audience = doc.add_paragraph()
    audience.paragraph_format.space_before = Pt(18)
    audience.paragraph_format.space_after = Pt(0)
    ar = audience.add_run("Hazırlanan kitle: Ürün, mühendislik, editoryal, AI kalite ve beta operasyon ekipleri")
    set_run_font(ar, size=9.5, color=MUTED, italic=True)

    doc.add_page_break()
    markdown_body(doc, lines[body_start:])

    props = doc.core_properties
    props.title = "LocalAkademi Ürün Stratejisi"
    props.subject = "Kontrollü beta ve v1.0 ürünleşme stratejisi"
    props.author = "LocalAkademi"
    props.keywords = "LocalAkademi, ürün stratejisi, kontrollü beta, eğitim, AI Mentor"
    props.comments = "LA-MP-02 v1.0"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
